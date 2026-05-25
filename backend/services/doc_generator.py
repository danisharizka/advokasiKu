from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

TEMPLATE_HEADER = {
    "surat_pengaduan": "SURAT PENGADUAN",
    "somasi": "SURAT SOMASI",
    "keberatan": "SURAT KEBERATAN",
    "permohonan": "SURAT PERMOHONAN",
}

def generate_docx(doc_data: dict) -> bytes:
    """
    Generate surat resmi .docx dari data terstruktur.
    Kembalikan bytes untuk langsung dikirim sebagai response.
    """
    doc = Document()

    # --- Page margins ---
    section = doc.sections[0]
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1.2)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.2)

    # --- Judul surat ---
    doc_type = doc_data.get("doc_type", "surat_pengaduan")
    title = doc.add_paragraph(TEMPLATE_HEADER.get(doc_type, "SURAT RESMI"))
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # --- Nomor dan tanggal ---
    today = datetime.now().strftime("%d %B %Y")
    info_table = doc.add_table(rows=2, cols=2)
    info_table.style = "Table Grid"
    info_table.cell(0, 0).text = "Nomor"
    info_table.cell(0, 1).text = f": ___/ADV/{datetime.now().strftime('%Y')}"
    info_table.cell(1, 0).text = "Tanggal"
    info_table.cell(1, 1).text = f": {today}"

    # Hapus border tabel
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for row in info_table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    doc.add_paragraph()

    # --- Kepada ---
    recipient = doc_data.get("recipient", "Yang Berwenang")
    p_to = doc.add_paragraph()
    p_to.add_run("Kepada Yth.").bold = True
    doc.add_paragraph(f"Bapak/Ibu {recipient}")
    doc.add_paragraph("di Tempat")

    doc.add_paragraph()

    # --- Perihal ---
    p_subj = doc.add_paragraph()
    p_subj.add_run("Perihal  : ").bold = True
    p_subj.add_run(doc_data.get("data", {}).get("perihal", "Pengaduan"))

    doc.add_paragraph()

    # --- Salam pembuka ---
    doc.add_paragraph("Dengan hormat,")
    doc.add_paragraph()

    # --- Isi surat ---
    uraian = doc_data.get("data", {}).get("uraian", "")
    if uraian:
        p_isi = doc.add_paragraph(uraian)
        p_isi.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph()

    # --- Dasar hukum ---
    pasal_list = doc_data.get("data", {}).get("pasal_relevan", [])
    if pasal_list:
        p_dh = doc.add_paragraph()
        p_dh.add_run("Adapun dasar hukum yang menjadi landasan pengaduan ini adalah:").bold = True
        for pasal in pasal_list:
            doc.add_paragraph(f"• {pasal}", style="List Bullet")

    doc.add_paragraph()

    # --- Tuntutan ---
    tuntutan = doc_data.get("data", {}).get("tuntutan", "")
    if tuntutan:
        p_tut = doc.add_paragraph()
        p_tut.add_run("Maka dengan ini kami memohon agar:").bold = True
        doc.add_paragraph(tuntutan)

    doc.add_paragraph()

    # --- Penutup ---
    doc.add_paragraph(
        "Demikian surat ini kami sampaikan. Atas perhatian dan tindak lanjutnya, "
        "kami ucapkan terima kasih."
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # --- Tanda tangan ---
    ttd_table = doc.add_table(rows=5, cols=2)
    ttd_table.cell(0, 0).text = "Hormat kami,"
    ttd_table.cell(0, 1).text = f"..............., {today}"
    ttd_table.cell(1, 0).text = ""
    ttd_table.cell(1, 1).text = "Pengirim,"
    ttd_table.cell(2, 0).text = ""
    ttd_table.cell(3, 0).text = ""
    ttd_table.cell(4, 0).text = "( ___________________________ )"
    ttd_table.cell(4, 1).text = "( ___________________________ )"

    # --- Save ke bytes ---
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
